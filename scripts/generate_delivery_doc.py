"""Gera o documento de ENTREGA do Trabalho 1 (Word, formatado) em docs/."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# --- Helpers de formatação --------------------------------------------------

def add_hr(p):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def shade_cell(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_table_header(row, fill="1F3864", text_color=(255, 255, 255)):
    for cell in row.cells:
        shade_cell(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(*text_color)
                run.font.size = Pt(10)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tbl_pr.append(borders)


# --- Documento --------------------------------------------------------------

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Estilos base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Cabeçalho institucional ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("JARVIS Acadêmico")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Documento de Entrega — Trabalho 1")
r.italic = True
r.font.size = Pt(13)
add_hr(sub)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.add_run(
    "Disciplina de Inteligência Artificial\n"
    "Universidade Federal de Mato Grosso do Sul (UFMS) · Faculdade de Computação (FACOM)\n"
    f"Acadêmico: Felipe Sá · Ano: 2026 · Data: {date.today().isoformat()}"
)
doc.add_paragraph("")

# --- 1. Contextualização ---
doc.add_heading("1. Contextualização do trabalho", level=1)
doc.add_paragraph(
    "Este documento acompanha a entrega do Trabalho 1 da disciplina de "
    "Inteligência Artificial. O sistema desenvolvido, denominado JARVIS "
    "Acadêmico, é um assistente pessoal inteligente para estudantes "
    "universitários, que combina três técnicas modernas de IA exigidas pelo "
    "enunciado:"
)
for item in [
    "Retrieval-Augmented Generation (RAG) sobre materiais de estudo do próprio "
    "estudante (PDFs, textos, anotações);",
    "Tool Calling com decisão autônoma do modelo, em um agente que executa, "
    "encadeia e audita ferramentas durante a conversa;",
    "Integração com o modelo Gemma 12B servido pelo endpoint OpenAI-"
    "compatible da LIA UFMS.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_paragraph(
    "O escopo do Trabalho 1 cobre as três funcionalidades obrigatórias da "
    "primeira entrega:"
)
funcs = [
    ("3.1 — Consulta a materiais de estudo (RAG)",
     "ingestão de PDF/TXT/MD em chunks, embeddings multilíngues, recuperação "
     "vetorial e geração de respostas com citações."),
    ("3.2 — Agenda acadêmica",
     "criação e consulta de eventos (aulas, provas, trabalhos) com data/hora, "
     "tipo, local e descrição."),
    ("3.3 — Lista de tarefas",
     "criação, listagem, conclusão e exclusão de tarefas com prazo e "
     "prioridade."),
]
for nome, desc in funcs:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(nome + " — ").bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Como diferencial (item 14 do enunciado), o sistema entrega ainda: (i) "
    "interface web moderna em modo escuro com sidebar colapsável e calendário "
    "mensal unificando eventos e tarefas, (ii) auditoria completa de toda "
    "chamada de ferramenta em SQLite, consultável pela própria UI, e (iii) "
    "persistência de conversas anteriores com restauração via barra lateral."
)

# --- 2. Tecnologias utilizadas ---
doc.add_heading("2. Tecnologias utilizadas", level=1)
doc.add_paragraph(
    "A stack foi escolhida com prioridade para bibliotecas pequenas, código "
    "explícito e total controle sobre cada camada (sem frameworks que "
    "implementam RAG ou agent loops automaticamente)."
)
tech_groups = [
    ("Linguagem e ambiente", "Python 3.12; gerenciador uv (Astral); venv isolada"),
    ("Modelo de linguagem (LLM)",
     "Gemma 12B servido em endpoint OpenAI-compatible da LIA UFMS; cliente "
     "AsyncOpenAI; biblioteca tenacity para retry com backoff exponencial em "
     "falhas transitórias (timeout, 5xx, 429)"),
    ("Embeddings e RAG",
     "sentence-transformers com o modelo intfloat/multilingual-e5-small "
     "(PT/EN); biblioteca sqlite-vec para busca vetorial nativa no SQLite; "
     "pdfplumber para extração de texto de PDFs; Recursive Character Splitter "
     "próprio (800 chars / 150 de overlap)"),
    ("Persistência",
     "SQLite (módulo sqlite3 da stdlib) em modo WAL com foreign_keys ativo; "
     "migrations forward-only versionadas via PRAGMA user_version; Pydantic v2 "
     "para validação e modelagem; pydantic-settings para configuração via .env"),
    ("Interface gráfica",
     "NiceGUI 3.x (servidor Python + Quasar/Vue no browser); fonte Inter via "
     "Google Fonts; locale Quasar PT-BR injetado para calendários e "
     "datepickers; CSS customizado próprio (sem dependência adicional de "
     "Bootstrap ou Material UI)"),
    ("Observabilidade",
     "loguru para logging estruturado (console colorizado + arquivo rotativo "
     "diário); tabela SQLite tool_call_logs registrando entrada, saída, "
     "status e duração de cada chamada de ferramenta"),
    ("Qualidade",
     "pytest + pytest-asyncio para testes unitários e de integração; ruff "
     "para lint e formatação; mypy (modo não-estrito) para checagem de tipos"),
]
table_tech = doc.add_table(rows=1, cols=2)
table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table_tech.rows[0].cells
hdr[0].text = "Camada"
hdr[1].text = "Tecnologias"
style_table_header(table_tech.rows[0])
for camada, tec in tech_groups:
    row = table_tech.add_row().cells
    row[0].text = camada
    row[1].text = tec
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_table_borders(table_tech)

# --- 3. Arquitetura de pastas ---
doc.add_heading("3. Arquitetura de pastas", level=1)
doc.add_paragraph(
    "A organização do código segue camadas com regras explícitas de "
    "importação, prevenindo ciclos e mantendo cada responsabilidade isolada."
)

p = doc.add_paragraph()
p.add_run(
    "src/core/ → src/llm/ → src/tools/ → src/domain/ + src/rag/ → src/ui/"
).italic = True

doc.add_paragraph(
    "Cada camada pode importar apenas das camadas inferiores; src/core/ não "
    "importa nada interno (apenas a stdlib do Python e bibliotecas externas)."
)

doc.add_paragraph("A estrutura completa do repositório é:")

tree = """
.
├── README.md
├── pyproject.toml
├── uv.lock
├── .env.example
├── .gitignore
│
├── src/
│   ├── main.py                     Entry point (sobe a UI)
│   ├── core/                       Infraestrutura (DB, config, logging, health)
│   │   └── migrations/             001_initial · 002_rag · 003_chat · 004_calendar_view
│   ├── llm/                        Cliente Gemma + agent loop
│   ├── rag/                        Ingest, chunk, embed, retrieve, pipeline
│   ├── domain/                     Regras de negócio
│   │   ├── agenda/                 Eventos
│   │   ├── tasks/                  Tarefas
│   │   ├── chat/                   Sessões e mensagens persistidas
│   │   └── calendar_view/          Leitura unificada (VIEW SQL)
│   ├── tools/                      8 ferramentas + registry
│   └── ui/                         Interface NiceGUI
│       ├── app.py                  Página /
│       ├── theme.py                CSS global + locale Quasar PT-BR
│       ├── state.py                Estado singleton
│       ├── components/             Sidebar, chat, calendário, etc.
│       └── dialogs/                Materiais, calendário, tarefas, auditoria
│
├── tests/
│   ├── unit/                       Testes unitários
│   └── integration/                Integração contra SQLite real
│
├── data/                           Dataset RAG (PDFs, .txt, .md)
├── docs/                           Relatórios de entrega (este e dataset)
├── img/                            Logos institucionais
└── logs/                           Saída do loguru (rotação diária)
"""
code_p = doc.add_paragraph()
run = code_p.add_run(tree.strip("\n"))
run.font.name = "Consolas"
run.font.size = Pt(9)

doc.add_paragraph(
    "Os números atuais do projeto: 99 arquivos versionados, ~10.200 linhas de "
    "código e configuração, 4 migrations SQL, 8 ferramentas, 7 tabelas + 1 "
    "VIEW unificada, 107 testes automatizados (104 passando, 3 marcados como "
    "skip por exigirem o endpoint LLM real)."
)

# --- 4. Localização do trabalho no GitHub ---
doc.add_heading("4. Localização do trabalho no GitHub", level=1)
doc.add_paragraph(
    "Todo o código-fonte, configuração e documentação pública estão "
    "versionados no repositório:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("https://github.com/felipegomessa/jarvis")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph(
    "O repositório segue o padrão de \"clonar e rodar\" descrito no README, "
    "com instruções para Windows, macOS e Linux. O setup completo (ambiente, "
    "dependências, variáveis de ambiente, primeira execução) está coberto na "
    "seção Instalação do README."
)
p = doc.add_paragraph()
p.add_run("Comandos resumidos: ").bold = True
p.add_run(
    "git clone https://github.com/felipegomessa/jarvis.git · "
    "uv sync --extra dev · "
    "preencher .env com JARVIS_LLM_API_KEY · "
    "python -m src.main · "
    "abrir http://127.0.0.1:8080."
)

# --- 5. Localização do dataset ---
doc.add_heading("5. Localização do dataset", level=1)
doc.add_paragraph(
    "O dataset utilizado para a funcionalidade de RAG (10 documentos "
    "acadêmicos, ~91 MB no total) está disponível em pasta compartilhada do "
    "Google Drive:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "https://drive.google.com/drive/folders/1X1FOAFXexwKrtveNSZ4zYnFfl55NPgpd"
)
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph(
    "O dataset está organizado em três subpastas — Livros, Artigos e Material "
    "de Aula — e cobre fundamentos clássicos de IA (Russell & Norvig, "
    "Mitchell), deep learning (Goodfellow et al.), tópicos específicos "
    "(regressão logística, KNN, TF-IDF) e conteúdo metodológico sobre RAG "
    "(\"Adaptive Chunking\"). A documentação completa do dataset, contendo "
    "origem, tipo, categoria, estratégia de chunking aplicada documento a "
    "documento, impacto no RAG e limitações observadas, está no arquivo "
    "complementar docs/Dataset_e_Chunking.docx (entregue junto com este)."
)

# --- 6. Uso do Claude Code no desenvolvimento ---
doc.add_heading("6. Uso do Claude Code no desenvolvimento", level=1)
doc.add_paragraph(
    "Em conformidade com o item 11 (\"Uso de IA para desenvolvimento\") e o "
    "item 13 (\"Itens permitidos\") do enunciado, declaramos de forma "
    "transparente que a ferramenta Claude Code (CLI de assistência ao "
    "desenvolvimento de software) foi utilizada durante a construção deste "
    "projeto. O Claude Code está explicitamente listado entre as ferramentas "
    "permitidas pelo enunciado."
)
doc.add_paragraph(
    "O uso seguiu o princípio de manter o controle pleno do código pelo "
    "acadêmico: cada arquivo, decisão arquitetural e linha de código pode ser "
    "explicada por Felipe Sá, conforme exigido pelos itens 10 e 11 do "
    "enunciado. O Claude Code foi empregado nos seguintes pontos específicos:"
)

pontos = [
    ("Processo de especificação (SDD).",
     "Entrevistas guiadas com o acadêmico produziram, para cada funcionalidade, "
     "uma tripla requirements.md + design.md + tasks.md, com critérios de "
     "aceitação verificáveis antes da escrita do código. Foram produzidas 7 "
     "specs versionadas (000 a 006), uma para cada área do sistema."),
    ("Geração de boilerplate sob revisão.",
     "Modelos Pydantic, repositórios de acesso a SQLite, migrations .sql, "
     "definições de configuração via pydantic-settings e fixtures de teste "
     "foram gerados a partir das specs aprovadas e revisados pelo acadêmico "
     "antes da integração."),
    ("Implementação do pipeline de RAG.",
     "Extração de texto com pdfplumber, Recursive Character Splitter próprio, "
     "embeddings com sentence-transformers, indexação em sqlite-vec e "
     "recuperação por similaridade — código revisado em cada etapa."),
    ("Implementação do agent loop de Tool Calling.",
     "Encadeamento de até 6 iterações, parser de JSON estrito com 1 reparo "
     "automático em caso de saída malformada, logging completo em "
     "tool_call_logs e tratamento explícito de erros por categoria."),
    ("Construção da interface gráfica.",
     "Reescrita completa da UI inicial (modelo de abas básico) para o layout "
     "ChatGPT-style atual (sidebar colapsável de 260px↔60px, chat central com "
     "input em pill, calendário unificado de eventos+tarefas, dialogs "
     "modais). A versão final reúne componentes próprios em "
     "src/ui/components/ e src/ui/dialogs/."),
    ("Detecção e correção de bugs sutis.",
     "Casos identificados incluíram: o bug de COMMIT implícito em "
     "sqlite3.Connection.executescript() que quebrava atomicidade de "
     "migrations; ordem incorreta dos handlers de exceção do SDK openai (que "
     "fazia RateLimitError ser engolido por APIStatusError, perdendo retries); "
     "comportamento do componente ui.upload (que exigia clique adicional "
     "implícito), corrigido com auto_upload=True. Cada correção foi validada "
     "por testes específicos."),
    ("Suíte de testes automatizados.",
     "107 testes (104 passando + 3 marcados como skip por exigirem o endpoint "
     "real do LLM): testes unitários para chunking, prompt RAG, lógica de "
     "cores, saudações, modelos Pydantic; testes de integração contra um "
     "SQLite temporário real para repos, migrations e agent loop. Lint via "
     "ruff sem warnings."),
    ("Documentação técnica.",
     "README atual do repositório, este documento de entrega, o documento "
     "complementar sobre o dataset (docs/Dataset_e_Chunking.docx) e o "
     "roteiro do vídeo demonstrativo foram redigidos com apoio do Claude "
     "Code, e revisados pelo acadêmico antes do uso."),
    ("Operações de versionamento e publicação.",
     "Inicialização do repositório Git local com proteção das credenciais "
     "(.env) via .gitignore, configuração do autor local apenas neste "
     "repositório, commit inicial limpo (99 arquivos, sem dados sensíveis) e "
     "push para o GitHub do acadêmico."),
]
for titulo, desc in pontos:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(titulo + " ").bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Em nenhum momento foram utilizadas ferramentas proibidas pelo item 13 do "
    "enunciado (lovable, Bolt.new, v0, Pythagora AI, Replit Ghostwriter, "
    "Flowise, Retool AI, Superblocks AI, BuildShip, Glide, Bubble, Appsmith, "
    "ToolJet, entre outras). Todo o sistema — RAG, integração com LLM e tool "
    "calling — foi implementado explicitamente em código próprio, conforme "
    "exige o item 10 do enunciado."
)

# --- 7. Conformidade com a entrega ---
doc.add_heading("7. Conformidade com a entrega exigida (item 12)", level=1)
doc.add_paragraph(
    "A tabela a seguir mapeia cada exigência da seção 12 do enunciado ao "
    "respectivo entregável."
)
reqs = [
    ("12.1 — Código em repositório GitHub",
     "https://github.com/felipegomessa/jarvis"),
    ("12.1 — README com instruções",
     "README.md na raiz do repositório, com setup, execução, testes e fluxo de uso"),
    ("12.1 — Lista de IAs utilizadas",
     "Seção 6 deste documento (Claude Code, com pontos de uso descritos)"),
    ("12.2 — Dataset com mínimo de 10 documentos",
     "10 documentos PDF no Google Drive (link na Seção 5)"),
    ("12.2 — Documentação do dataset",
     "docs/Dataset_e_Chunking.docx (origem, tipo, categoria, chunking, impacto no RAG, limitações)"),
    ("12.3 — Vídeo de até 3 minutos",
     "Roteiro elaborado; gravação será entregue separadamente"),
]
table_req = doc.add_table(rows=1, cols=2)
table_req.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table_req.rows[0].cells
hdr[0].text = "Exigência"
hdr[1].text = "Onde está atendida"
style_table_header(table_req.rows[0])
for req, where in reqs:
    row = table_req.add_row().cells
    row[0].text = req
    row[1].text = where
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_table_borders(table_req)

# --- 8. Autoria ---
doc.add_heading("8. Autoria", level=1)
doc.add_paragraph(
    "Trabalho desenvolvido pelo acadêmico Felipe Sá, para a disciplina de "
    "Inteligência Artificial da Universidade Federal de Mato Grosso do Sul "
    "(UFMS), Faculdade de Computação (FACOM), 2026."
)

# --- Salvar ---
out_path = Path("docs") / "Entrega_Trabalho1.docx"
out_path.parent.mkdir(exist_ok=True)
doc.save(out_path)
print(f"OK: documento salvo em {out_path.resolve()}")
