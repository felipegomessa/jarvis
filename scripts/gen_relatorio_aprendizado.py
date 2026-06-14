"""Gera o relatório Word das Melhorias de Aprendizado (Trabalho 2) — RF-007.10/T-007.15.

Reaproveita o `docs/Entrega_Trabalho1.docx` como TEMPLATE: herda fonte (Calibri),
estilos de título (Heading 1=14pt, Heading 2=13pt), código (Courier) e o tema do
documento, garantindo a MESMA padronização de formatação. O conteúdo antigo é
removido e o relatório é escrito por cima, salvando em outro arquivo.

Uso: python scripts/gen_relatorio_aprendizado.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "Entrega_Trabalho1.docx"
OUTPUT = ROOT / "docs" / "Relatorio_Melhorias_Aprendizado.docx"


def _clear_body(doc: Document) -> None:
    """Remove todo o conteúdo do corpo, preservando as propriedades de seção."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _has_style(doc: Document, name: str) -> bool:
    return any(s.name == name for s in doc.styles)


def build(doc: Document) -> None:
    code_style = "macro" if _has_style(doc, "macro") else "Normal"
    bullet_style = "List Bullet" if _has_style(doc, "List Bullet") else "Normal"

    def heading(text: str, level: int = 1) -> None:
        doc.add_paragraph(text, style=f"Heading {level}")

    def para(text: str = "") -> None:
        doc.add_paragraph(text)

    def rich(parts: list[tuple[str, bool]]) -> None:
        p = doc.add_paragraph()
        for txt, bold in parts:
            run = p.add_run(txt)
            run.bold = bold

    def bullet(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph(style=bullet_style)
        run = p.add_run(text)
        run.bold = bold

    def code(lines: str) -> None:
        doc.add_paragraph(lines, style=code_style)

    # ---------------------------------------------------------------- Capa
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("JARVIS Acadêmico")
    r.bold = True
    r.font.size = Pt(22)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Relatório de Melhorias de Aprendizado — Trabalho 2")
    r.bold = True
    r.font.size = Pt(13)

    for _ in range(2):
        para()
    for line, center in [
        ("Disciplina de Inteligência Artificial", True),
        ("Universidade Federal de Mato Grosso do Sul (UFMS)", True),
        ("Faculdade de Computação (FACOM)", True),
        ("Acadêmicos: Felipe Sá e Eduardo Sá", True),
        ("Junho de 2026", True),
    ]:
        p = doc.add_paragraph(line)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ----------------------------------------------------------- 0. Sumário
    heading("Sumário", 1)
    for item in [
        "1. Introdução e objetivo",
        "2. Atendimento ao requisito (mapeamento)",
        "3. Visão geral da arquitetura",
        "4. Funcionalidade 1 — Prova eletrônica (interativa)",
        "5. Funcionalidade 2 — Identificação de dificuldades e plano de estudos",
        "6. Integração e tool calling",
        "7. Evidências de funcionamento",
        "8. Qualidade de engenharia",
        "9. Limitações conhecidas",
        "10. Conclusão",
        "11. Apêndice — IAs utilizadas e referências",
    ]:
        bullet(item)
    doc.add_page_break()

    # ------------------------------------------------------ 1. Introdução
    heading("1. Introdução e objetivo", 1)
    para(
        "Este relatório documenta as funcionalidades de Melhorias de Aprendizado "
        "implementadas no JARVIS Acadêmico, em atendimento ao requisito obrigatório "
        "do Trabalho 2. O enunciado exige a implementação de pelo menos duas "
        "funcionalidades voltadas ao aprendizado do usuário, das quais ao menos uma "
        "deve ser interativa (o sistema pergunta e avalia)."
    )
    para("Foram entregues duas funcionalidades integradas:")
    bullet(
        "Prova eletrônica: a partir dos materiais de estudo carregados, o sistema "
        "gera uma prova com questões de múltipla escolha e dissertativas; o aluno "
        "responde online e o sistema corrige e atribui uma nota de 0 a 10. Esta é a "
        "funcionalidade interativa."
    )
    bullet(
        "Identificação de dificuldades e plano de estudos: a partir do desempenho na "
        "prova, o sistema identifica os tópicos de maior dificuldade e monta um plano "
        "de estudos didático, consultando a agenda do aluno para distribuir as "
        "sessões em horários livres."
    )
    para(
        "Ambas reutilizam os três pilares do sistema — RAG (recuperação sobre os "
        "materiais), o modelo de linguagem (LLM) e o tool calling — desenvolvidos no "
        "Trabalho 1."
    )

    # ------------------------------------------------ 2. Mapeamento ao requisito
    heading("2. Atendimento ao requisito (mapeamento)", 1)
    para(
        "A tabela a seguir mapeia cada exigência do enunciado à forma como foi "
        "atendida nesta entrega."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(("Exigência do enunciado", "Como foi atendida", "Interativa?")):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(txt)
        run.bold = True
    linhas = [
        (
            "Pelo menos 2 funcionalidades de aprendizado",
            "Prova eletrônica (geração + correção + nota) e Identificação de "
            "dificuldades + plano de estudos.",
            "—",
        ),
        (
            "Pelo menos 1 funcionalidade interativa (pergunta e avalia)",
            "A prova eletrônica: o sistema gera as questões, o aluno responde e o "
            "sistema corrige e atribui nota 0 a 10 com feedback.",
            "Sim",
        ),
        (
            "Apoiar o aprendizado e integrar múltiplas fontes",
            "Questões e recomendações são aterradas no material (RAG); o plano "
            "integra material, desempenho e agenda do aluno.",
            "—",
        ),
    ]
    for a, b, c in linhas:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c
    para()

    # ------------------------------------------------ 3. Arquitetura
    heading("3. Visão geral da arquitetura", 1)
    para(
        "As funcionalidades foram implementadas respeitando a separação de camadas "
        "do projeto. Para combinar LLM, RAG e domínio sem violar as regras de "
        "dependência, foi criada uma nova camada de orquestração, registrada na "
        "decisão de arquitetura D-030."
    )
    bullet(
        "src/domain/learning/ — modelos e repositório (provas, questões, tentativas, "
        "respostas). Camada pura, sem dependência de LLM.",
    )
    bullet(
        "src/learning/ — orquestração: generator (gera as questões), grader (corrige), "
        "coach (dificuldades e plano) e service (fluxo de alto nível).",
    )
    bullet(
        "src/llm/client.py e json_utils.py — cliente LLM padrão (injeção sem acoplar "
        "a interface) e parsing tolerante de JSON do modelo.",
    )
    bullet(
        "src/rag/retrieve.get_document_chunks — leitura de um documento inteiro, em "
        "ordem (base da geração de provas).",
    )
    bullet(
        "src/tools/tool_learning.py — cinco ferramentas decididas pela LLM.",
    )
    bullet(
        "src/ui/dialogs/exam_dialog.py — interface interativa da prova (menu '+').",
    )
    para(
        "A persistência usa a migration 005 (versão 5 do banco), que adiciona as "
        "tabelas quizzes, quiz_documents, quiz_questions, quiz_attempts e "
        "quiz_answers, mantendo o histórico de provas e tentativas."
    )

    # ------------------------------------------------ 4. Funcionalidade 1
    heading("4. Funcionalidade 1 — Prova eletrônica (interativa)", 1)

    heading("4.1 Descrição e objetivo", 2)
    para(
        "Permite que o aluno teste seus conhecimentos a partir do próprio material "
        "de estudo. O sistema gera uma prova, recebe as respostas e devolve uma nota "
        "de 0 a 10 com gabarito comentado — caracterizando o comportamento "
        "interativo exigido (pergunta e avalia)."
    )

    heading("4.2 Fluxo do usuário", 2)
    bullet("Passo 1 — Configurar: o aluno seleciona um ou mais materiais, define o "
           "número de questões de múltipla escolha e dissertativas e escolhe o idioma "
           "da prova (português, padrão, ou o idioma original do material).")
    bullet("Passo 2 — Responder: as questões são exibidas (rádio para múltipla "
           "escolha, caixa de texto para dissertativas).")
    bullet("Passo 3 — Corrigir: ao enviar, o sistema corrige automaticamente.")
    bullet("Passo 4 — Resultado: nota 0 a 10 em destaque, a resposta do aluno ao "
           "lado da resposta correta, e feedback por questão.")

    heading("4.3 Fundamentação técnica", 2)
    para(
        "Geração aterrada no material: o gerador lê os trechos (chunks) dos "
        "documentos selecionados e instrui a LLM a criar questões apenas com base "
        "neles, registrando em cada questão o trecho de origem. Isso reduz a "
        "alucinação e torna as questões rastreáveis ao material."
    )
    para(
        "Correção: as questões de múltipla escolha são corrigidas de forma "
        "determinística (comparação do índice escolhido com o gabarito). As "
        "dissertativas são avaliadas por um LLM-juiz que compara a resposta do aluno "
        "com uma rubrica e o trecho-fonte, produzindo uma nota parcial (de 0 a 1) e "
        "um feedback. A nota da dissertativa é apresentada como sugestão, com o "
        "gabarito visível, por transparência."
    )
    para("A nota final é calculada como a soma dos pontos obtidos dividida pela soma "
         "dos pontos possíveis, multiplicada por 10.")

    heading("4.4 Robustez (qualidade de engenharia)", 2)
    bullet("Reparo de JSON: se o modelo retorna um JSON inválido ou incompleto, o "
           "sistema faz uma tentativa de reparo via re-prompt antes de desistir.")
    bullet("Validação flexível de alternativas: questões de múltipla escolha aceitam "
           "de 2 a 6 alternativas (4 como alvo), evitando descartar boas questões "
           "quando o modelo varia o número de opções.")
    bullet("Tratamento de erros: falhas são reportadas ao usuário de forma amigável "
           "e registradas em log, sem quebrar a interface.")

    # ------------------------------------------------ 5. Funcionalidade 2
    heading("5. Funcionalidade 2 — Identificação de dificuldades e plano de estudos", 1)

    heading("5.1 Identificação de dificuldades", 2)
    para(
        "Cada questão é associada a um tópico. Após a correção, o sistema agrega o "
        "desempenho por tópico (pontos obtidos sobre possíveis) e classifica como "
        "fracos os tópicos com aproveitamento abaixo de um limiar configurável."
    )

    heading("5.2 Plano de estudos didático e consciente da agenda", 2)
    para("Para os tópicos fracos, o sistema gera um plano de estudos com três "
         "características:")
    bullet("Didático: ações concretas e progressivas (reler uma seção, resumir com "
           "as próprias palavras, fazer exercícios, explicar em voz alta), citando o "
           "material de apoio.")
    bullet("Dimensionado pela profundidade: tópicos com menor aproveitamento recebem "
           "mais sessões e/ou mais tempo de estudo.")
    bullet("Consciente da agenda: o sistema consulta os compromissos do aluno nos "
           "próximos dias e distribui as sessões em dias e horários livres, evitando "
           "conflito com eventos já agendados.")
    para(
        "O plano é apresentado ao aluno e pode ser adicionado à agenda com um clique: "
        "as sessões com data e hora viram eventos no horário sugerido; itens sem "
        "horário viram tarefas. Essa integração entre material, desempenho e agenda "
        "antecipa a Funcionalidade 3.4 (planejamento de estudos) do enunciado."
    )

    # ------------------------------------------------ 6. Tool calling
    heading("6. Integração e tool calling", 1)
    para(
        "As funcionalidades também são acessíveis por linguagem natural no chat, pois "
        "foram expostas como ferramentas que o próprio modelo decide quando chamar:"
    )
    bullet("gerar_prova — cria uma prova a partir de materiais (com idioma e número "
           "de questões).")
    bullet("corrigir_prova — corrige uma tentativa e retorna a nota.")
    bullet("identificar_dificuldades — aponta os tópicos de maior dificuldade.")
    bullet("montar_plano_estudos — monta o plano de estudos.")
    bullet("ler_documento — lê um documento inteiro em ordem (útil para 'resuma o "
           "material X' ou 'me dê o índice').")
    para(
        "Toda chamada de ferramenta é registrada na tabela de auditoria "
        "(tool_call_logs), com entrada, saída, status e duração, atendendo ao "
        "requisito de logs de tool calling."
    )

    # ------------------------------------------------ 7. Evidências
    heading("7. Evidências de funcionamento", 1)
    para(
        "As capturas de tela a seguir comprovam o funcionamento das funcionalidades "
        "com o modelo e o material reais. (Inserir as imagens nos espaços indicados.)"
    )
    for legenda in [
        "[Inserir print: tela de configuração da prova — seleção de materiais, "
        "número de questões e idioma]",
        "[Inserir print: questões sendo respondidas (múltipla escolha e dissertativa)]",
        "[Inserir print: tela de resultado — nota 0 a 10, resposta do aluno ao lado "
        "da correta e feedback]",
        "[Inserir print: dificuldades por tópico e plano de estudos com dias/horários]",
        "[Inserir print: tab de Auditoria mostrando as chamadas de ferramenta "
        "registradas]",
    ]:
        p = doc.add_paragraph(legenda)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.italic = True

    # ------------------------------------------------ 8. Engenharia
    heading("8. Qualidade de engenharia", 1)
    bullet("Testes automatizados: a suíte do projeto passa com 167 casos aprovados e "
           "3 ignorados (testes de LLM ao vivo, opcionais). As funcionalidades "
           "adicionaram cerca de 40 novos testes cobrindo modelos, correção, parsing "
           "da prova, coach, leitura por documento, migration 005 e geração com LLM "
           "simulado.")
    bullet("Tratamento de erros explícito: JSON malformado, LLM indisponível, "
           "documento ilegível e resposta em branco são tratados com mensagens "
           "amigáveis e logs.")
    bullet("Separação de responsabilidades: domínio, orquestração, ferramentas e "
           "interface em camadas com regras de dependência explícitas.")
    bullet("Processo: a funcionalidade seguiu o processo de desenvolvimento dirigido "
           "por especificação (Spec 007), auditada por um agente revisor antes da "
           "implementação, e registrada na decisão de arquitetura D-030.")
    bullet("Padrão de código: análise estática (ruff) sem apontamentos e verificação "
           "de tipos (mypy).")

    # ------------------------------------------------ 9. Limitações
    heading("9. Limitações conhecidas", 1)
    bullet("A correção de questões dissertativas depende do julgamento da LLM; por "
           "isso a nota é apresentada como sugestão, acompanhada do gabarito e do "
           "feedback.")
    bullet("A qualidade das questões depende da qualidade e do idioma do material "
           "indexado; materiais com pouco texto extraível limitam a geração.")
    bullet("Documentos muito longos são amostrados por cota de trechos para caber no "
           "contexto do modelo.")
    bullet("O planejamento considera os compromissos cadastrados na agenda; janelas "
           "livres dependem de a agenda estar atualizada.")

    # ------------------------------------------------ 10. Conclusão
    heading("10. Conclusão", 1)
    para(
        "As duas funcionalidades atendem ao requisito obrigatório de Melhorias de "
        "Aprendizado, sendo a prova eletrônica a funcionalidade interativa exigida. "
        "A solução apoia efetivamente o estudo do aluno: avalia o conhecimento sobre "
        "o próprio material, identifica onde estão as dificuldades e propõe um plano "
        "de estudos realista, integrado à agenda. A implementação reaproveita e "
        "fortalece os pilares de RAG, LLM e tool calling do sistema."
    )

    # ------------------------------------------------ 11. Apêndice
    heading("11. Apêndice — IAs utilizadas e referências", 1)
    para("Ferramentas de IA utilizadas no desenvolvimento:")
    bullet("Claude Code (Anthropic) — desenvolvimento assistido, geração de código, "
           "revisão e auditoria de especificação.")
    para("Referências internas do projeto:")
    bullet("Especificação da funcionalidade: spec/007-learning/ (requirements, "
           "design, tasks, audit).")
    bullet("Decisão de arquitetura: D-030 (camada learning, cliente LLM padrão, "
           "leitura por documento) em decisions.md.")
    bullet("Modelo de linguagem: Qwen2.5-14B-Instruct via endpoint da LIA/UFMS "
           "(decisão D-028).")


def main() -> None:
    if not TEMPLATE.exists():
        raise SystemExit(f"template nao encontrado: {TEMPLATE}")
    doc = Document(str(TEMPLATE))
    _clear_body(doc)
    doc.core_properties.title = "Relatorio de Melhorias de Aprendizado - Trabalho 2"
    doc.core_properties.author = "Felipe Sa; Eduardo Sa"
    build(doc)
    doc.save(str(OUTPUT))
    print(f"OK: relatorio gerado em {OUTPUT}")


if __name__ == "__main__":
    main()
