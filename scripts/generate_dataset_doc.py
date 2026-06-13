"""Gera um relatório Word (.docx) sobre o dataset e a estratégia de chunking.

Lê o JSON produzido por analyze_dataset.py e produz um documento formatado em
docs/Dataset_e_Chunking.docx.
"""

from __future__ import annotations

import json
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Carrega dados ------------------------------------------------------------
STATS_PATH = Path(tempfile.gettempdir()) / "dataset_stats.json"
data: list[dict] = json.loads(STATS_PATH.read_text(encoding="utf-8"))

# Catálogo "humano" — origem inferida, tipo, observações por arquivo
CATALOG: dict[str, dict[str, str]] = {
    "Adaptive_Chunking_Optimizing_Chunking-Method_Selection_for_RAG.pdf": {
        "origem": "Repositório acadêmico (pré-print) — autoria Paulo R. de Moura Júnior, Jean Lelong, Annabelle Blangero",
        "tipo": "Artigo científico",
        "lingua": "Inglês",
        "tema": "Seleção adaptativa de métodos de chunking para RAG",
    },
    "The_Origins_of_Logistic_Regression.pdf": {
        "origem": "Artigo histórico/acadêmico (autoria não declarada no metadata; conteúdo clássico sobre regressão logística)",
        "tipo": "Artigo científico",
        "lingua": "Inglês",
        "tema": "História e fundamentos da regressão logística",
    },
    "Understanding TF-IDF (Term Frequency-Inverse Document Frequency) - GeeksforGeeks.pdf": {
        "origem": "Artigo web exportado para PDF — GeeksforGeeks",
        "tipo": "Artigo de divulgação (web)",
        "lingua": "Inglês",
        "tema": "Conceito de TF-IDF aplicado a recuperação de informação",
    },
    "art_of_ml.pdf": {
        "origem": "Livro digitalizado (metadados anônimos)",
        "tipo": "Livro",
        "lingua": "Inglês",
        "tema": "Aprendizado de máquina (geral)",
    },
    "Artificial_Intelligence_A_Modern_Approach,_Prentice_Hall_Series_in_Artificial_Intelligence.pdf": {
        "origem": "Stuart Russell & Peter Norvig — \"Artificial Intelligence: A Modern Approach\" (AIMA), 4ª ed.",
        "tipo": "Livro (clássico de referência da disciplina)",
        "lingua": "Inglês",
        "tema": "Inteligência Artificial — visão geral abrangente",
    },
    "ddidl.pdf": {
        "origem": "Livro digitalizado (metadados ausentes)",
        "tipo": "Livro",
        "lingua": "Inglês",
        "tema": "Conteúdo técnico de grande extensão (área de dados/sistemas)",
    },
    "Ian Goodfellow, Yoshua Bengio, Aaron Courville - Deep Learning (2017, MIT).pdf": {
        "origem": "Ian Goodfellow, Yoshua Bengio, Aaron Courville — \"Deep Learning\" (MIT Press, 2017)",
        "tipo": "Livro (referência canônica)",
        "lingua": "Inglês",
        "tema": "Deep Learning — fundamentos e aplicações",
    },
    "Text Book Machine LEarning.pdf": {
        "origem": "Tom M. Mitchell — \"Machine Learning\"",
        "tipo": "Livro (referência clássica)",
        "lingua": "Inglês",
        "tema": "Aprendizado de máquina — fundamentos",
    },
    "aula-KNN.pdf": {
        "origem": "Material didático universitário (autor: gbatista) — slides sobre KNN",
        "tipo": "Material de aula (slides)",
        "lingua": "Português",
        "tema": "Algoritmo K-Nearest Neighbors",
    },
    "lista1.pdf": {
        "origem": "Material didático universitário — Lista de exercícios 1",
        "tipo": "Material de aula (lista de exercícios)",
        "lingua": "Português",
        "tema": "Exercícios introdutórios da disciplina",
    },
}


# Helpers de formatação -----------------------------------------------------

def add_hr(p):
    """Adiciona uma linha horizontal sob o parágrafo."""
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def shade_cell(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_table_header(row, fill="1F3864", text_color=(255, 255, 255)):
    for cell in row.cells:
        shade_cell(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(*text_color)
                run.font.size = Pt(10)


def set_table_borders(table):
    """Aplica bordas simples a uma tabela inteira."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tbl_pr.append(borders)


# Constrói o documento -----------------------------------------------------

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Estilos base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Capa / Cabeçalho ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("JARVIS Acadêmico")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Dataset e Estratégia de Chunking — Relatório Técnico")
r.italic = True
r.font.size = Pt(13)
add_hr(sub)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.add_run(
    "Trabalho 1 — Disciplina de Inteligência Artificial\n"
    "Universidade Federal de Mato Grosso do Sul (UFMS) · Faculdade de Computação (FACOM) · 2026\n"
    f"Autor: Felipe Sá · Data: {date.today().isoformat()}"
)

doc.add_paragraph("")

# --- 1. Introdução ---
doc.add_heading("1. Introdução", level=1)
doc.add_paragraph(
    "Este relatório documenta o dataset construído para o sistema JARVIS Acadêmico, "
    "respondendo aos requisitos das seções 8 (Dataset) e 12.2 (Entrega) do enunciado "
    "do Trabalho 1. São apresentados:"
)
for item in [
    "a relação completa dos documentos disponíveis (nome, origem, tipo e categoria);",
    "a estratégia de chunking utilizada e como ela se aplica a cada documento;",
    "o impacto dessa estratégia sobre a etapa de recuperação (RAG);",
    "as limitações observadas no dataset.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

p = doc.add_paragraph()
p.add_run("Local físico do dataset: ").bold = True
p.add_run(
    "pasta compartilhada no Google Drive em "
    "https://drive.google.com/drive/folders/1X1FOAFXexwKrtveNSZ4zYnFfl55NPgpd "
    "(10 arquivos PDF, ~91 MB no total). Os arquivos podem ser ingeridos via a "
    "tela de upload da própria UI ou pelo comando de ingestão em lote da pasta "
    "/data."
)

# --- 2. Documentos do dataset ---
doc.add_heading("2. Documentos do dataset", level=1)
p = doc.add_paragraph()
p.add_run("Total de documentos: ").bold = True
p.add_run(f"{len(data)} (atende ao mínimo exigido pelo enunciado, item 8).")

cat_counts: dict[str, int] = {}
for d in data:
    cat_counts[d["category"]] = cat_counts.get(d["category"], 0) + 1
p = doc.add_paragraph()
p.add_run("Distribuição por categoria: ").bold = True
p.add_run(
    ", ".join(f"{k} ({v})" for k, v in sorted(cat_counts.items())) + "."
)

# Tabela principal
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Arquivo"
hdr[1].text = "Categoria"
hdr[2].text = "Tipo / Origem"
hdr[3].text = "Tema"
style_table_header(table.rows[0])

for d in data:
    meta = CATALOG.get(d["file"], {})
    row = table.add_row().cells
    row[0].text = d["file"]
    row[1].text = d["category"]
    row[2].text = f"{meta.get('tipo','-')}\n{meta.get('origem','-')}"
    row[3].text = meta.get("tema", "-")
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
set_table_borders(table)

# --- 3. Estratégia de chunking ---
doc.add_heading("3. Estratégia de chunking", level=1)

doc.add_heading("3.1 Método empregado", level=2)
doc.add_paragraph(
    "O JARVIS Acadêmico utiliza um Recursive Character Text Splitter implementado em "
    "src/rag/chunk.py. A divisão respeita uma hierarquia de separadores semânticos "
    "(quebras de parágrafo, quebras de linha, finais de sentença, espaços e, em "
    "último caso, caracteres individuais), priorizando manter unidades de significado "
    "íntegras dentro de cada chunk."
)
doc.add_paragraph(
    "Os parâmetros de operação, definidos como padrão do projeto (decisão D-006 — "
    "documentada internamente), são:"
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Tamanho-alvo de chunk: 800 caracteres.").bold = False
p = doc.add_paragraph(style="List Bullet")
p.add_run("Sobreposição entre chunks consecutivos: 150 caracteres.").bold = False
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Estratégia uniforme para todos os documentos — o mesmo splitter é aplicado a "
    "todo o dataset, sem regras por tipo de arquivo."
).bold = False

doc.add_paragraph(
    "A escolha do tamanho 800/150 considera: (i) a janela de contexto utilizada na "
    "geração final pela LLM Gemma 12B, que beneficia trechos auto-contidos; (ii) a "
    "granularidade necessária para o modelo de embeddings multilingual-e5-small, cujo "
    "comprimento máximo de entrada é de 512 tokens (~2 000 caracteres em inglês, "
    "~1 200 em português); e (iii) a sobreposição de 150 caracteres preserva "
    "continuidade entre fronteiras, reduzindo o risco de perder uma sentença "
    "dividida ao meio."
)

doc.add_heading("3.2 Pipeline de ingestão", level=2)
doc.add_paragraph(
    "Cada PDF do dataset passa pelas seguintes etapas, implementadas em "
    "src/rag/ingest.py:"
)
steps = [
    "Cálculo de hash SHA-256 do arquivo para deduplicação idempotente.",
    "Extração de texto página a página com pdfplumber.",
    "Aplicação do Recursive Character Splitter (800/150).",
    "Geração de embeddings densos (intfloat/multilingual-e5-small).",
    "Persistência atômica em SQLite (tabelas documents, chunks) e em sqlite-vec (chunk_vecs).",
]
for s in steps:
    p = doc.add_paragraph(s, style="List Number")

doc.add_heading("3.3 Aplicação do chunking por documento", level=2)
doc.add_paragraph(
    "A tabela a seguir mostra, para cada documento, o volume textual extraído e a "
    "quantidade de chunks gerados pela aplicação do splitter."
)

table2 = doc.add_table(rows=1, cols=6)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = "Arquivo"
hdr2[1].text = "Páginas"
hdr2[2].text = "Caracteres"
hdr2[3].text = "Palavras"
hdr2[4].text = "Chunks"
hdr2[5].text = "Média chars/chunk"
style_table_header(table2.rows[0])

total_chunks = 0
total_chars = 0
for d in data:
    row = table2.add_row().cells
    row[0].text = d["file"]
    row[1].text = f"{d['pdf_pages']}"
    row[2].text = f"{d['text_chars']:,}".replace(",", ".")
    row[3].text = f"{d['text_words']:,}".replace(",", ".")
    row[4].text = f"{d['chunks_generated']:,}".replace(",", ".")
    row[5].text = f"{d['avg_chunk_chars']:.1f}"
    total_chunks += d["chunks_generated"]
    total_chars += d["text_chars"]
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
# Linha de totais
total_row = table2.add_row().cells
total_row[0].text = "TOTAL"
total_row[1].text = "—"
total_row[2].text = f"{total_chars:,}".replace(",", ".")
total_row[3].text = "—"
total_row[4].text = f"{total_chunks:,}".replace(",", ".")
total_row[5].text = "—"
for c in total_row:
    shade_cell(c, "D9E2F3")
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
set_table_borders(table2)

doc.add_paragraph(
    f"O dataset gera aproximadamente {total_chunks:,} chunks indexados, "
    f"a partir de {total_chars:,} caracteres extraídos no total."
    .replace(",", ".")
)

# --- 4. Impacto no RAG ---
doc.add_heading("4. Impacto no RAG", level=1)

doc.add_heading("4.1 Aspectos positivos", level=2)
positives = [
    (
        "Cobertura temática ampla. ",
        "Os documentos cobrem fundamentos clássicos (Russell & Norvig, Mitchell), "
        "deep learning (Goodfellow et al.), tópicos específicos (KNN, regressão "
        "logística, TF-IDF) e conteúdo metodológico sobre RAG. Isso permite "
        "responder perguntas tanto introdutórias quanto aprofundadas.",
    ),
    (
        "Granularidade adequada à LLM. ",
        "O tamanho de 800 caracteres por chunk costuma manter uma ideia auto-"
        "contida (um parágrafo curto ou meia página), o que reduz a chance de "
        "trechos recuperados ficarem ambíguos ou cortados em um conceito-chave.",
    ),
    (
        "Sobreposição protege fronteiras. ",
        "Os 150 caracteres de overlap evitam que uma sentença relevante seja "
        "dividida em dois chunks distintos, preservando continuidade lexical e "
        "semântica entre vizinhos.",
    ),
    (
        "Compatibilidade com embeddings multilíngues. ",
        "O modelo multilingual-e5-small lida com inglês e português; o dataset "
        "tem materiais nas duas línguas, e o chunker não exige adaptação por "
        "idioma.",
    ),
]
for bold, text in positives:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(bold).bold = True
    p.add_run(text)

doc.add_heading("4.2 Aspectos que afetam negativamente", level=2)
negatives = [
    (
        "Volumes muito desiguais. ",
        "Há grande disparidade entre os documentos: o maior livro gera ~3.300 "
        "chunks, enquanto o menor artigo gera apenas 11. Documentos volumosos "
        "podem dominar os resultados de busca por similaridade, especialmente em "
        "consultas genéricas.",
    ),
    (
        "Extração ruidosa em PDFs complexos. ",
        "Vários PDFs apresentaram avisos do pdfplumber (\"Could not get FontBBox\"), "
        "típicos de fontes incorporadas mal especificadas. Isso pode produzir "
        "texto com glifos ausentes, ordem incorreta ou colunas embaralhadas, "
        "comprometendo embeddings desses trechos.",
    ),
    (
        "Chunking insensível a estrutura. ",
        "O Recursive Character Splitter não conhece títulos de seções, tabelas ou "
        "figuras. Quebras podem ocorrer no meio de listas, fórmulas, blocos de "
        "código ou legendas, fragmentando informações que idealmente deveriam "
        "permanecer juntas.",
    ),
    (
        "Mistura de idiomas no índice. ",
        "Como o dataset é predominantemente em inglês (8/10) mas as perguntas do "
        "usuário tendem a ser em português, a recuperação por similaridade "
        "depende fortemente da qualidade do alinhamento cross-lingual do "
        "multilingual-e5-small — qualidade que cai em terminologias muito "
        "específicas.",
    ),
    (
        "PDFs escaneados ou com OCR de baixa qualidade. ",
        "Alguns documentos digitalizados produzem texto fragmentado. Quando isso "
        "acontece, o chunking opera sobre texto ruim e os embeddings carregam "
        "esse ruído.",
    ),
]
for bold, text in negatives:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(bold).bold = True
    p.add_run(text)

doc.add_heading("4.3 Síntese", level=2)
doc.add_paragraph(
    "A estratégia 800/150 com Recursive Character Splitter atende ao escopo do "
    "trabalho: é simples de explicar, reproduzível, e gera chunks com tamanho "
    "compatível tanto com o modelo de embeddings quanto com a LLM. Em troca, "
    "perde-se sensibilidade à estrutura semântica de cada documento. Para "
    "trabalhos futuros, métodos de chunking estruturados (por seção, por sentença "
    "ou adaptativos — como propõe o artigo \"Adaptive Chunking\" presente no "
    "próprio dataset) podem mitigar parte das limitações listadas em 4.2."
)

# --- 5. Limitações do dataset ---
doc.add_heading("5. Limitações do dataset", level=1)
limitations = [
    (
        "Tamanho mínimo respeitado, sem grande margem. ",
        "São exatamente 10 documentos, o mínimo exigido pelo enunciado (item 8). "
        "Pequenas perdas (extração inviável, conteúdo descartado) reduziriam o "
        "número efetivo de documentos abaixo do mínimo.",
    ),
    (
        "Distribuição linguística desbalanceada. ",
        "8 documentos em inglês e 2 em português. Consultas em português sobre "
        "tópicos cobertos apenas em livros em inglês exigem que o embedding e a "
        "LLM façam tradução implícita, com perda de precisão semântica.",
    ),
    (
        "Metadados ausentes ou incorretos. ",
        "Vários PDFs não trazem título nem autor no metadata (por exemplo, "
        "ddidl.pdf, Ian Goodfellow Deep Learning, The Origins of Logistic "
        "Regression). Um deles (aula-KNN) tem título de metadata divergente do "
        "conteúdo (\"Ordenação em Memória Interna\"), provavelmente reuso de "
        "template. Isso prejudica buscas por filtros (autor/título) e a "
        "rastreabilidade nas citações.",
    ),
    (
        "Origem incerta em alguns arquivos. ",
        "art_of_ml.pdf traz o autor \"Sality [BЯ]\" no metadata — referência a "
        "fonte não acadêmica. Documentos com proveniência incerta levantam "
        "questões de licenciamento e de qualidade da edição.",
    ),
    (
        "Volume textual muito heterogêneo. ",
        "Documentos variam de 7 mil a mais de 2 milhões de caracteres. Mesmo com "
        "normalização por similaridade, livros longos dominam estatisticamente o "
        "espaço de busca; tópicos curtos (TF-IDF, lista de exercícios) ficam "
        "subrepresentados.",
    ),
    (
        "Cobertura desigual de subáreas. ",
        "Há concentração em fundamentos clássicos de ML/IA e deep learning. "
        "Temas específicos solicitados em provas e listas (ex.: avaliação de "
        "modelos, validação cruzada, ética em IA) podem aparecer apenas como "
        "menções marginais.",
    ),
    (
        "Ausência de notas e anotações pessoais. ",
        "O dataset é constituído exclusivamente de PDFs publicados; não inclui "
        "anotações próprias do estudante (caderno, resumos), que são fontes "
        "valiosas no contexto pessoal do assistente.",
    ),
    (
        "PDFs com fontes problemáticas. ",
        "Alvos como o artigo The_Origins_of_Logistic_Regression apresentaram "
        "chunks/página atipicamente alto (~24 chunks/página), indicando "
        "extração possivelmente comprometida — caracteres duplicados ou "
        "delimitadores residuais aumentam artificialmente o volume textual.",
    ),
]
for bold, text in limitations:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(bold).bold = True
    p.add_run(text)

# --- 6. Conformidade com o enunciado ---
doc.add_heading("6. Conformidade com o enunciado", level=1)
doc.add_paragraph(
    "A tabela abaixo mapeia cada exigência das seções 8 e 12.2 do enunciado ao "
    "ponto deste documento que a atende."
)

reqs = [
    ("Item 8 — mínimo de 10 documentos", "Seção 2 (10 documentos listados)"),
    ("Item 8 — conteúdo acadêmico", "Seção 2 (livros canônicos, artigos científicos e material de aula)"),
    ("Item 8 — origem dos dados", "Seção 2 (tabela com origem por arquivo)"),
    ("Item 8 — tipo de conteúdo", "Seção 2 (coluna 'Tipo/Origem')"),
    ("Item 8 — limitações", "Seção 5 (oito limitações listadas)"),
    ("Item 8 — estratégia de chunking", "Seção 3"),
    ("Item 8 — impacto no RAG", "Seção 4"),
    ("Item 8 — pasta /data ou link", "Link Google Drive informado na Seção 1"),
    ("Item 12.2 — documentação completa", "Este documento (docs/Dataset_e_Chunking.docx)"),
]
table3 = doc.add_table(rows=1, cols=2)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
hdr3[0].text = "Exigência do enunciado"
hdr3[1].text = "Onde está atendida"
style_table_header(table3.rows[0])
for req, where in reqs:
    row = table3.add_row().cells
    row[0].text = req
    row[1].text = where
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_table_borders(table3)

# --- 7. Rodapé / autoria ---
doc.add_heading("7. Autoria", level=1)
doc.add_paragraph(
    "Documento elaborado pelo acadêmico Felipe Sá, para a disciplina de "
    "Inteligência Artificial da Universidade Federal de Mato Grosso do Sul (UFMS), "
    "Faculdade de Computação (FACOM), 2026."
)

# --- Salvar ---
out_path = Path("docs") / "Dataset_e_Chunking.docx"
out_path.parent.mkdir(exist_ok=True)
doc.save(out_path)
print(f"OK: documento salvo em {out_path.resolve()}")
