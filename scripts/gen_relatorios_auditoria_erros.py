"""Gera 2 relatórios Word: Auditoria do Enunciado + Análise de Erros (Trabalho 2).

Reaproveita `docs/Entrega_Trabalho1.docx` como TEMPLATE (mesma fonte Calibri,
estilos de título, tema) — igual ao gerador do relatório de aprendizado.

Uso: python scripts/gen_relatorios_auditoria_erros.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "Entrega_Trabalho1.docx"
OUT_AUDIT = ROOT / "docs" / "Relatorio_Auditoria_Enunciado.docx"
OUT_ERROS = ROOT / "docs" / "Relatorio_Analise_Erros.docx"


# --------------------------------------------------------------------- helpers
def _clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _has_style(doc: Document, name: str) -> bool:
    return any(s.name == name for s in doc.styles)


class W:
    """Pequeno wrapper de escrita com os estilos do template."""

    def __init__(self, doc: Document) -> None:
        self.doc = doc
        self.bullet = "List Bullet" if _has_style(doc, "List Bullet") else "Normal"

    def h(self, text: str, level: int = 1) -> None:
        self.doc.add_paragraph(text, style=f"Heading {level}")

    def p(self, text: str = "") -> None:
        self.doc.add_paragraph(text)

    def b(self, text: str, bold: bool = False) -> None:
        par = self.doc.add_paragraph(style=self.bullet)
        run = par.add_run(text)
        run.bold = bold

    def cover(self, subtitle: str) -> None:
        t = self.doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("JARVIS Acadêmico")
        r.bold = True
        r.font.size = Pt(22)
        s = self.doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run(subtitle)
        r.bold = True
        r.font.size = Pt(13)
        self.p()
        self.p()
        for line in [
            "Disciplina de Inteligência Artificial",
            "Universidade Federal de Mato Grosso do Sul (UFMS)",
            "Faculdade de Computação (FACOM)",
            "Acadêmicos: Felipe Sá e Eduardo Sá",
            "Junho de 2026",
        ]:
            par = self.doc.add_paragraph(line)
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_page_break()

    def table(self, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        tb = self.doc.add_table(rows=1, cols=len(headers))
        tb.style = "Table Grid"
        for i, htxt in enumerate(headers):
            cell = tb.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(htxt)
            run.bold = True
        for row in rows:
            cells = tb.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        if widths:
            for r_ in tb.rows:
                for i, wd in enumerate(widths):
                    r_.cells[i].width = Pt(wd)
        self.p()


# --------------------------------------------------------------- AUDITORIA doc
OK = "Correto"
PARC = "Parcialmente correto"
NOK = "Incorreto"


def build_auditoria(doc: Document) -> None:
    w = W(doc)
    w.cover("Relatório de Auditoria de Conformidade ao Enunciado")

    w.h("1. Objetivo", 1)
    w.p(
        "Este relatório audita a conformidade do projeto JARVIS Acadêmico aos "
        "requisitos do enunciado do trabalho, considerando as entregas do Trabalho 1 "
        "(funcionalidades 3.1 a 3.3) e do Trabalho 2 (3.4, melhorias de aprendizado, "
        "avaliação e análise de erros). Cada item é formulado como uma pergunta "
        "objetiva e classificado como Correto, Parcialmente correto ou Incorreto, "
        "com justificativa."
    )
    w.p("Legenda das classificações:")
    w.b("Correto: o requisito é plenamente atendido.")
    w.b("Parcialmente correto: atendido em parte, com ressalva justificada.")
    w.b("Incorreto: não atendido ou ainda pendente.")

    w.h("2. Funcionalidades obrigatórias", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["3.1 — O sistema consulta materiais de estudo via RAG (carrega "
             "documentos, divide em chunks, gera embeddings, recupera trechos e "
             "gera respostas baseadas neles)?", OK,
             "RAG próprio: pdfplumber + chunking recursivo + multilingual-e5-small "
             "+ sqlite-vec; respostas com citação [Doc N] e bloco de fontes."],
            ["3.2 — O sistema responde consultas à agenda acadêmica (o que tenho "
             "hoje, aulas da semana, prova amanhã)?", OK,
             "Domínio de agenda com eventos e consultas por período; tools "
             "consultar_agenda/consultar_calendario."],
            ["3.3 — O sistema permite adicionar, listar e concluir tarefas?", OK,
             "Domínio de tarefas (CRUD) + tools adicionar_tarefa, listar_tarefas, "
             "concluir_tarefa (por id ou título)."],
            ["3.4 — O sistema monta um plano de estudos combinando agenda, tarefas "
             "e materiais?", PARC,
             "Entregue como plano de estudos didático a partir do desempenho na "
             "prova: integra materiais (RAG) + agenda (evita horários ocupados) e "
             "cria tarefas/eventos. Um assistente geral de priorização diária "
             "('o que priorizar hoje') ainda pode ser ampliado."],
        ],
    )

    w.h("3. Tool calling", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["O sistema implementa pelo menos 5 ferramentas?", OK,
             "15 ferramentas registradas (agenda, tarefas, RAG, calendário, "
             "prova/aprendizado, leitura de documento)."],
            ["A decisão de chamar a ferramenta é feita pela LLM (não lógica fixa)?",
             OK, "Agent loop próprio: a LLM responde em JSON escolhendo a tool; o "
             "loop executa e injeta a observação até a resposta final."],
            ["Cada chamada registra logs com ferramenta, entrada e saída?", OK,
             "Tabela tool_call_logs com I/O em JSON, status e duração; consultável "
             "na própria UI (aba Auditoria)."],
        ],
    )

    w.h("4. Melhorias de aprendizado", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["Há pelo menos 2 funcionalidades voltadas ao aprendizado?", OK,
             "Prova eletrônica (geração + correção + nota) e identificação de "
             "dificuldades + plano de estudos."],
            ["Pelo menos uma é interativa (o sistema pergunta e avalia)?", OK,
             "A prova eletrônica: gera questões, o aluno responde, o sistema "
             "corrige (MC determinística + dissertativa por LLM-juiz) e dá nota 0-10."],
        ],
    )

    w.h("5. Avaliação e análise de erros", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["O grupo avaliou o sistema com pelo menos 10 perguntas (pergunta, "
             "documentos recuperados, resposta e classificação)?", NOK,
             "Pendente: a avaliação sistemática de >=10 perguntas ainda será "
             "executada e registrada (artefato complementar a este pacote)."],
            ["O grupo identificou pelo menos 3 falhas (tipo, causa, solução)?", OK,
             "Entregue no Relatório de Análise de Erros: 4 falhas (2 corrigidas, "
             "2 em aberto com solução proposta)."],
        ],
    )

    w.h("6. Dataset", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["O dataset tem pelo menos 10 documentos acadêmicos?", OK,
             "11 documentos (artigos, material de aula e livros-texto de IA/ML)."],
            ["Está em /data no repositório ou link externo?", OK,
             "Versionado em data/ (Artigos, Material de Aula, Livros)."],
            ["A documentação cobre origem, tipo e limitações?", OK,
             "data/README.md descreve procedência, categorias e limitações."],
            ["A estratégia de chunking e o impacto no RAG estão explicados?", OK,
             "data/README.md detalha o Recursive Character Splitter (~800/150) e o "
             "impacto observado (9.502 chunks, validação de recuperação)."],
        ],
    )

    w.h("7. Qualidade de engenharia", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["O código demonstra organização e separação de responsabilidades?", OK,
             "Camadas (core, domain, rag, llm, learning, tools, ui) com regras de "
             "dependência explícitas (CLAUDE.md §4.1)."],
            ["Há testes básicos?", OK,
             "167 testes (unit + integration) verdes; smoke de LLM opcional."],
            ["Há tratamento de erros e logs?", OK,
             "Política de erros (CLAUDE.md §8); loguru com arquivo rotativo; erros "
             "tratados, nunca silenciosos."],
            ["O grupo usou IA para revisão/sugestão/bugs e consegue explicar o "
             "código?", OK,
             "Claude Code (desenvolvimento, revisão, auditoria de spec); processo "
             "SDD com ADRs garante rastreabilidade e explicabilidade."],
        ],
    )

    w.h("8. Tecnologias e restrições", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["O grupo implementou explicitamente RAG, integração com LLM e tool "
             "calling (sem framework que gere o sistema)?", OK,
             "Implementação própria, sem LangChain/afins; bibliotecas pequenas."],
            ["O sistema usa a LLM obrigatória (Gemma 12B) fornecida pelo professor?",
             PARC,
             "O projeto foi construído sobre o Gemma 12B no endpoint da LIA/UFMS; "
             "a LIA aposentou o Gemma e passou a servir o Qwen2.5-14B. O cliente é "
             "OpenAI-compatível e a troca foi apenas de configuração (decisão "
             "D-028). Recomenda-se confirmar com o professor."],
            ["O desenvolvimento evitou ferramentas proibidas (geradores de sistema "
             "completo)?", OK,
             "Usou apenas ferramentas permitidas (Claude Code)."],
        ],
    )

    w.h("9. Entrega", 1)
    w.table(
        ["Item (pergunta)", "Avaliação", "Justificativa"],
        [
            ["Há repositório GitHub com o código?", OK,
             "github.com/felipegomessa/jarvis (branch main)."],
            ["Há README com instruções de execução?", OK,
             "README com setup (uv sync, .env, execução) e índice pré-construído."],
            ["O README lista as IAs utilizadas?", OK,
             "Claude Code listado; demais ferramentas registradas."],
            ["O dataset foi entregue conforme exigido?", OK,
             "Sim (ver seção 6)."],
            ["Há vídeo de até 3 minutos com arquitetura e sistema funcionando?",
             PARC,
             "O vídeo do Trabalho 1 (3.1/3.2/3.3) foi entregue; o vídeo do Trabalho "
             "2 (prova, plano de estudos) ainda será gravado."],
        ],
    )

    w.h("10. Síntese", 1)
    w.p(
        "A maioria dos requisitos está plenamente atendida. Pontos de atenção: "
        "(a) a avaliação com >=10 perguntas ainda será executada; (b) o vídeo do "
        "Trabalho 2 será gravado; (c) a LLM efetivamente em uso é o Qwen2.5-14B, "
        "pois a LIA/UFMS aposentou o Gemma 12B — situação a confirmar com o "
        "professor. A funcionalidade 3.4 foi entregue de forma parcial, via plano "
        "de estudos integrado à agenda."
    )


# ----------------------------------------------------------------- ERROS doc
def build_analise_erros(doc: Document) -> None:
    w = W(doc)
    w.cover("Relatório de Análise de Erros do Sistema")

    w.h("1. Introdução e metodologia", 1)
    w.p(
        "Este relatório atende ao requisito de análise de erros do Trabalho 2: "
        "identificar pelo menos 3 falhas do sistema, cada uma com tipo, causa e "
        "possível solução. As falhas foram identificadas por inspeção do código, "
        "auditoria do índice vetorial e observação do comportamento em uso real "
        "(geração de respostas e de provas). Foram encontradas quatro falhas: duas "
        "já corrigidas durante o desenvolvimento e duas mantidas em aberto com "
        "solução proposta. A avaliação sistemática com >=10 perguntas é um artefato "
        "complementar a ser executado."
    )

    w.h("2. Resumo das falhas", 1)
    w.table(
        ["#", "Falha", "Tipo", "Status"],
        [
            ["1", "Extração de PDF ilegível indexada como lixo", "Recuperação / "
             "Ingestão", "Corrigida"],
            ["2", "Trechos irrelevantes enviados como contexto", "Recuperação / "
             "Relevância", "Em aberto"],
            ["3", "Resposta sem aterramento obrigatório no material", "Geração",
             "Em aberto"],
            ["4", "Impossibilidade de ler um documento específico", "Ambiguidade / "
             "Cobertura", "Corrigida"],
        ],
    )

    w.h("3. Falha 1 — Extração de PDF ilegível indexada como lixo", 1)
    w.b("Tipo: recuperação / ingestão (extração de texto).", True)
    w.b("Causa: o documento The Origins of Logistic Regression usa fontes sem mapa "
        "de caracteres (sem ToUnicode). pdfplumber, pdfminer e PyMuPDF (todos "
        "testados) extraem apenas marcadores (cid:N), sem texto real. A ingestão só "
        "verificava se o texto era vazio — o lixo, por não ser vazio, era embeddado "
        "e poluía o índice (100% dos 383 chunks com (cid:, 0% de palavras reais), "
        "violando a própria política de tratamento de erros do projeto.")
    w.b("Evidência: consulta ao índice mostrava 383 chunks ilegíveis; pedir o "
        "índice do documento retornava trechos de outros documentos.")
    w.b("Possível solução (aplicada): guarda de qualidade na ingestão "
        "(real_word_ratio) que rejeita texto com menos de 25% de palavras reais, "
        "registrando aviso em log e pulando o arquivo (decisão D-029). O documento "
        "deixou de poluir o índice. Recuperá-lo exigiria OCR.")
    w.b("Status: CORRIGIDA.")

    w.h("4. Falha 2 — Trechos irrelevantes enviados como contexto", 1)
    w.b("Tipo: recuperação / relevância.", True)
    w.b("Causa: a função de recuperação marca 'sem contexto relevante' apenas com "
        "base na distância do MELHOR trecho, mas retorna todos os top-k trechos "
        "para a LLM, inclusive os irrelevantes (acima do limiar). O limiar funciona "
        "como porteiro do conjunto, não filtra trecho a trecho.")
    w.b("Evidência: em buscas cuja melhor correspondência é fraca, trechos de "
        "documentos não relacionados acompanham o contexto enviado ao modelo, "
        "aumentando o risco de respostas imprecisas.")
    w.b("Possível solução: filtrar individualmente cada trecho pela sua distância "
        "(descartar os acima do limiar) antes de compor o contexto; opcionalmente, "
        "aplicar re-ranking. Mitigado hoje pela instrução de aterramento no prompt.")
    w.b("Status: EM ABERTO (solução proposta).")

    w.h("5. Falha 3 — Resposta sem aterramento obrigatório no material", 1)
    w.b("Tipo: geração.", True)
    w.b("Causa: o prompt do sistema PEDE que a LLM use a busca em materiais para "
        "perguntas conceituais e cite as fontes, mas nada OBRIGA. O modelo pode "
        "responder a partir do próprio conhecimento, sem consultar o material — "
        "produzindo respostas plausíveis, porém não fundamentadas no acervo (e "
        "potencialmente incorretas ou desatualizadas).")
    w.b("Evidência: por ser uma decisão do modelo, há perguntas conceituais "
        "respondidas sem chamada à ferramenta de RAG.")
    w.b("Possível solução: forçar a recuperação antes de responder perguntas "
        "conceituais (pipeline que sempre recupera) e/ou validar que a resposta "
        "final cite [Doc N]; caso contrário, reexecutar com contexto recuperado.")
    w.b("Status: EM ABERTO (solução proposta).")

    w.h("6. Falha 4 — Impossibilidade de ler um documento específico", 1)
    w.b("Tipo: ambiguidade / cobertura de recuperação.", True)
    w.b("Causa: a busca em materiais fazia recuperação semântica GLOBAL, sem escopo "
        "por documento. Pedidos como 'leia o documento X', 'me dê o índice' ou "
        "'resuma o material X' eram inviáveis: alguns trechos soltos não "
        "reconstroem a estrutura de um documento, e a busca trazia trechos de "
        "outros documentos.")
    w.b("Evidência: ao pedir o índice de um documento específico, o sistema "
        "retornava trechos de documentos diferentes e recusava responder.")
    w.b("Possível solução (aplicada): função de leitura por documento "
        "(get_document_chunks, em ordem) e ferramenta ler_documento, criadas junto "
        "da funcionalidade de provas (Spec 007). Resolve leitura/índice/resumo de "
        "um documento específico.")
    w.b("Status: CORRIGIDA.")

    w.h("7. Falhas adicionais identificadas (acervo de melhorias)", 1)
    w.table(
        ["Falha", "Tipo", "Solução proposta", "Status"],
        [
            ["Trechos quase duplicados ocupam vagas do top-k (overlap sem dedup)",
             "Recuperação", "Remover duplicatas/quase-duplicatas antes de compor o "
             "contexto", "Em aberto"],
            ["Healthcheck do LLM só no boot (não detecta queda em uso)",
             "Robustez", "Verificação periódica + aviso na UI", "Em aberto"],
            ["Parâmetro top_k da configuração ignorado (valor fixo na tool)",
             "Configuração", "Usar o valor da configuração", "Em aberto"],
        ],
    )

    w.h("8. Conclusão", 1)
    w.p(
        "Foram identificadas quatro falhas reais, cobrindo tipos distintos "
        "(ingestão/recuperação, relevância, geração e ambiguidade de cobertura), "
        "o que supera o mínimo de três exigido. Duas falhas foram corrigidas "
        "durante o desenvolvimento (a indexação de PDF ilegível e a leitura por "
        "documento) e duas permanecem em aberto com solução proposta (filtragem "
        "de relevância por trecho e aterramento obrigatório). A análise demonstra "
        "uma postura crítica sobre o comportamento do sistema e orienta as "
        "próximas melhorias de qualidade do RAG."
    )


# --------------------------------------------------------------------- runner
def _gen(out: Path, builder) -> None:  # type: ignore[no-untyped-def]
    doc = Document(str(TEMPLATE))
    _clear_body(doc)
    builder(doc)
    doc.save(str(out))
    print(f"OK: {out.name}")


def main() -> None:
    if not TEMPLATE.exists():
        raise SystemExit(f"template nao encontrado: {TEMPLATE}")
    _gen(OUT_AUDIT, build_auditoria)
    _gen(OUT_ERROS, build_analise_erros)


if __name__ == "__main__":
    main()
